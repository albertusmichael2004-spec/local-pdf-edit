from pathlib import Path

from docx import Document

from backend.services.convert_to_pdf.jpg_to_text_to_pdf.docx_exporter import export_docx
from backend.services.convert_to_pdf.jpg_to_text_to_pdf.text_cleanup import clean_ocr_line
from backend.services.shared.ocr.models import OCRPage, OCRWord


def _page(text: str, page_number: int) -> OCRPage:
    words = [
        OCRWord(token, 95, 80 + index * 100, 100, 90, 30, 1, 1, 1)
        for index, token in enumerate(text.split())
    ]
    return OCRPage(
        Path(f"page-{page_number}.jpg"), 1000, 1400, 1000, 1400,
        words, 95, 3, "test",
    )


def test_indonesian_cleanup_is_conservative():
    assert clean_ocr_line("Jang berbicara dengan katakata") == "Yang berbicara dengan kata-kata"
    assert clean_ocr_line("la mengatakan demikian") == "Ia mengatakan demikian"
    assert clean_ocr_line("Gereja San Damiano.Dalam batin") == "Gereja San Damiano. Dalam batin"
    assert clean_ocr_line('"Sesudah la mengatakan, maka kiea') == '"Sesudah Ia mengatakan, maka kita'
    assert clean_ocr_line("Allah akan mengump") == "Allah akan mengumpulkan"


def test_cleanup_blocks_embedded_tesseract_metadata():
    leaked = "Bapa\n5\t1\t8\t1\t4\t6\t1538\t1372\t13\t19\t83.4\t,"
    assert clean_ocr_line(leaked) == "Bapa"


def test_docx_export_keeps_pages_editable_without_metadata(tmp_path: Path):
    output = tmp_path / "ocr.docx"
    export_docx([_page("Halaman pertama", 1), _page("Halaman kedua", 2)], output)
    document = Document(output)
    visible = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Halaman pertama" in visible
    assert "Halaman kedua" in visible
    assert "\t1\t" not in visible
    assert sum("w:type=\"page\"" in paragraph._p.xml for paragraph in document.paragraphs) == 1
